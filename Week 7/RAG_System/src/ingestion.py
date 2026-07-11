"""
ingestion.py
------------
Document Ingestion Module.

Accepts custom text inputs: PDFs, raw .txt files, .docx files, and
domain-specific JSON exports (e.g. Hugging Face dataset corpora such as
vectara/open_ragbench, whose corpus/*.json files store {"title", "sections":[...]}).

Responsibility: load raw files from a directory (or a single file) and
return a list of Document objects: {id, source, text}.
"""

import os
import json
from dataclasses import dataclass, field
from typing import List

SUPPORTED_EXTENSIONS = (".pdf", ".txt", ".md", ".docx", ".json")


@dataclass
class Document:
    doc_id: str
    source: str
    text: str
    metadata: dict = field(default_factory=dict)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_pdf(path: str) -> str:
    import fitz  # PyMuPDF
    text = []
    with fitz.open(path) as pdf:
        for page_num in range(pdf.page_count):
            page = pdf.load_page(page_num)
            text.append(page.get_text("text"))
    return "\n".join(text)


def _read_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


def _read_json(path: str) -> str:
    """
    Handles two shapes:
    1. A plain JSON string/list of strings.
    2. Open-RAG-Benchmark style corpus documents:
       {"title": ..., "abstract": ..., "sections": [{"text": ...}, ...]}
    """
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        data = json.load(f)

    if isinstance(data, str):
        return data
    if isinstance(data, list):
        return "\n".join(str(x) for x in data)
    if isinstance(data, dict):
        parts = []
        if "title" in data:
            parts.append(str(data["title"]))
        if "abstract" in data:
            parts.append(str(data["abstract"]))
        if "sections" in data and isinstance(data["sections"], list):
            for sec in data["sections"]:
                if isinstance(sec, dict) and "text" in sec:
                    parts.append(str(sec["text"]))
                else:
                    parts.append(str(sec))
        if not parts:
            # fallback: flatten all string values
            parts = [str(v) for v in data.values() if isinstance(v, (str, int, float))]
        return "\n".join(parts)
    return ""


READERS = {
    ".pdf": _read_pdf,
    ".txt": _read_txt,
    ".md": _read_txt,
    ".docx": _read_docx,
    ".json": _read_json,
}


def load_documents(path: str) -> List[Document]:
    """
    Load documents from a file or a directory of files.
    Unsupported extensions are silently skipped.
    """
    docs: List[Document] = []

    if os.path.isfile(path):
        files = [path]
    else:
        files = [
            os.path.join(path, fname)
            for fname in sorted(os.listdir(path))
            if fname.lower().endswith(SUPPORTED_EXTENSIONS)
        ]

    for fpath in files:
        ext = os.path.splitext(fpath)[1].lower()
        reader = READERS.get(ext)
        if reader is None:
            continue
        try:
            text = reader(fpath)
        except Exception as e:
            print(f"[ingestion] Failed to read {fpath}: {e}")
            continue
        text = text.strip()
        if not text:
            continue
        doc_id = os.path.splitext(os.path.basename(fpath))[0]
        docs.append(Document(doc_id=doc_id, source=fpath, text=text))

    return docs


if __name__ == "__main__":
    import sys
    d = load_documents(sys.argv[1] if len(sys.argv) > 1 else "data/sample_docs")
    for doc in d:
        print(doc.doc_id, len(doc.text), "chars")
